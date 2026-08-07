#!/usr/bin/env python3
"""DOCX для проверки системы — из текста настоящих законов, а не из выдумки.

Зачем это вообще нужно. Продукт принимает три формата: .pdf, .docx и .txt. В
системе на момент написания скрипта лежало семь документов — шесть .txt и один
.pdf, — то есть ветка разбора DOCX не проверялась НИ РАЗУ. Целый парсер жил на
доверии.

Почему собираем, а не скачиваем. Государственные сайты Таджикистана отдают
старый .doc (Composite Document File, Word 97), а не .docx: на mmk.tj рядом с
каждым PDF лежит именно `.doc`. Продукт .doc не принимает, конвертера в
окружении нет (ни libreoffice, ни pandoc, ни antiword), и ставить пакет на
полтора гигабайта ради тестовых файлов — плохой размен.

Поэтому текст берётся из НАСТОЯЩЕГО закона (его PDF), а контейнер собирается
python-docx. Проверяется ровно то, что нужно проверить: читает ли система
настоящий .docx с настоящим содержимым. И это сказано прямо, а не выдано за
скачанные документы: в манифесте у каждого файла стоит origin=converted и адрес
исходного PDF.

Заголовки расставляются по структуре закона («Моддаи N» — статья), а не
подряд одним абзацем: разбиение на фрагменты у продукта опирается на структуру
документа, и файл из одного абзаца проверил бы её мимо.

    ./venv/bin/python make_test_docx.py                 # из всех PDF корпуса
    ./venv/bin/python make_test_docx.py --limit 8
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
from pathlib import Path

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document  # noqa: E402

BACKEND_ROOT = Path(__file__).resolve().parent
CORPUS_ROOT = BACKEND_ROOT / "data" / "test_corpus"
LAWS_DIR = CORPUS_ROOT / "laws"
DOCX_DIR = CORPUS_ROOT / "docx"

# Строка, с которой начинается статья таджикского закона. По ней и ставится
# заголовок: «Моддаи 12. Ҳуқуқи шаҳрвандон» — это статья 12.
ARTICLE = re.compile(r"^\s*(Моддаи|Статья)\s+\d+", re.IGNORECASE)

# Абзац короче этого — обрывок колонтитула или номер страницы, а не текст.
MIN_PARAGRAPH = 40


def say(message: str = "") -> None:
    print(message, flush=True)


def blocks_of(pdf_path: Path) -> list[str]:
    """Абзацы PDF — тем же разбором, что и у продукта.

    Своего чтения PDF здесь нет намеренно: если продукт что-то читает неверно,
    тестовый файл должен унаследовать ту же неверность, иначе .docx и .pdf
    одного закона окажутся разными документами, и сравнить их будет нельзя.
    """
    from app.services.document_service import DocumentService

    blocks = DocumentService.extract_blocks(str(pdf_path), ".pdf")
    return [b.text.strip() for b in blocks if b.text and b.text.strip()]


def build(paragraphs: list[str], title: str, target: Path) -> int:
    document = Document()
    document.add_heading(title[:200] or target.stem, level=0)
    written = 0
    for text in paragraphs:
        if len(text) < MIN_PARAGRAPH and not ARTICLE.match(text):
            continue
        if ARTICLE.match(text):
            document.add_heading(text[:120], level=1)
        else:
            document.add_paragraph(text)
        written += 1
    document.save(target)
    return written


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    parser.add_argument("--laws", default=str(LAWS_DIR))
    parser.add_argument("--out", default=str(DOCX_DIR))
    parser.add_argument("--limit", type=int, default=None)
    args = parser.parse_args()

    laws = Path(args.laws)
    if not laws.is_dir():
        raise SystemExit(f"каталога {laws} нет — сначала соберите законы")

    described: dict[str, dict] = {}
    manifest_path = laws / "manifest.json"
    if manifest_path.is_file():
        raw = json.loads(manifest_path.read_text(encoding="utf-8"))
        rows = raw if isinstance(raw, list) else raw.get("files", [])
        described = {str(row.get("file")): row for row in rows if isinstance(row, dict)}

    pdfs = sorted(p for p in laws.glob("*.pdf"))
    if args.limit:
        pdfs = pdfs[: args.limit]
    if not pdfs:
        raise SystemExit(f"в {laws} нет ни одного PDF")

    out = Path(args.out)
    out.mkdir(parents=True, exist_ok=True)
    entries = []
    for pdf in pdfs:
        row = described.get(pdf.name, {})
        target = out / (pdf.stem + ".docx")
        try:
            paragraphs = blocks_of(pdf)
        except Exception as exc:  # разбор чужого PDF — место, где ломается всё
            say(f"  ! {pdf.name}: не разобрался ({exc})")
            continue
        if len(paragraphs) < 5:
            say(f"  ! {pdf.name}: всего {len(paragraphs)} абзацев, пропускаю")
            continue
        written = build(paragraphs, row.get("title", ""), target)
        entries.append(
            {
                "file": target.name,
                "title": row.get("title", ""),
                "subject": row.get("subject", ""),
                "origin": "converted",
                "source_pdf": pdf.name,
                "url": row.get("url", ""),
                "paragraphs": written,
                "size_bytes": target.stat().st_size,
            }
        )
        say(f"  {target.name}: {written} абзацев, {target.stat().st_size // 1024} КБ")

    (out / "manifest.json").write_text(
        json.dumps(entries, ensure_ascii=False, indent=1), encoding="utf-8"
    )
    say(f"\nсобрано {len(entries)} файлов .docx в {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
